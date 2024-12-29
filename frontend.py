import gradio as gr
import requests
from docx import Document

def extract_pdf_content(file):
    """Send a PDF file to the backend and return the extracted content."""
    if not file:
        return "No file provided. Please upload a PDF.", None

    # Define the API endpoint
    url = "https://file-extractor-api.onrender.com/extract"

    try:
        # Send the file to the backend
        with open(file.name, "rb") as f:
            response = requests.post(url, files={"file": f})

        # Check for errors
        if response.status_code != 200:
            return f"Error: {response.json().get('error', 'Unknown error')}", None

        # Extract content
        content = response.json().get("content", "No content extracted.")

        # Create a Word document
        doc = Document()
        doc.add_heading("Extracted Content", level=1)
        doc.add_paragraph(content)

        # Save the Word document
        output_file = "extracted_content.docx"
        doc.save(output_file)

        return content, output_file
    except Exception as e:
        return f"An error occurred: {str(e)}", None

# Define the Gradio interface
with gr.Blocks() as demo:
    gr.Markdown("## Content Extractor")
    gr.Markdown("Upload a PDF or Image file to extract its text content. Optionally, download the content as a Word document.")

    file_input = gr.File(label="Upload PDF")
    output_text = gr.Textbox(label="Extracted Content", lines=20, interactive=False)
    download_link = gr.File(label="Download Word File")

    extract_button = gr.Button("Extract Content")

    extract_button.click(
        extract_pdf_content, 
        inputs=[file_input], 
        outputs=[output_text, download_link]
    )

# Launch the Gradio app
demo.launch()
